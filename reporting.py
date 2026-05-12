# -*- coding: utf-8 -*-
from __future__ import annotations

from io import BytesIO
from typing import Dict

import pandas as pd
from docx import Document

from src import tabella_sintesi


def _fmt(value: object) -> str:
    if isinstance(value, (int, float)):
        return f"{value:.2f}"
    return str(value)


def _add_table(document: Document, df: pd.DataFrame) -> None:
    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for idx, col in enumerate(df.columns):
        table.rows[0].cells[idx].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, col in enumerate(df.columns):
            cells[idx].text = _fmt(row[col])


def create_word_report(sheets: Dict[str, pd.DataFrame], results: Dict[str, pd.DataFrame]) -> bytes:
    document = Document()
    document.add_heading("Relazione tecnica Telaio2D", 0)
    document.add_paragraph(
        "Analisi statica lineare di telaio piano. Le unita sono coerenti con gli input del modello."
    )

    document.add_heading("1. Sintesi", level=1)
    _add_table(document, tabella_sintesi(results))

    document.add_heading("2. Nodi", level=1)
    _add_table(document, sheets["nodes"])

    document.add_heading("3. Elementi", level=1)
    _add_table(document, sheets["elements"])

    document.add_heading("4. Risultati nodali", level=1)
    _add_table(document, results["results_nodal"])

    document.add_heading("5. Risultati elementi", level=1)
    _add_table(document, results["results_elements"])

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
