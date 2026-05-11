# -*- coding: utf-8 -*-
import io
from docx import Document
from docx.shared import Inches, Pt
import matplotlib.pyplot as plt
import pandas as pd
from datetime import datetime

def add_df_to_doc(doc, df: pd.DataFrame, font_size: int = 9):
    """Aggiunge un DataFrame pandas a un documento docx."""
    t = doc.add_table(df.shape[0] + 1, df.shape[1])
    t.style = 'Table Grid'
    for j, col_name in enumerate(df.columns):
        t.cell(0, j).text = col_name
    for i, row in enumerate(df.itertuples(), start=1):
        for j, val in enumerate(row[1:]):
            cell = t.cell(i, j)
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)

def plot_risultati_mpl(risultati, data_key, title):
    """Genera un grafico dei risultati con Matplotlib."""
    fig, ax = plt.subplots(figsize=(5, 8))
    ax.plot(risultati[data_key], -risultati['profondita_m'])
    ax.set_title(title)
    ax.set_xlabel(f"Valore [{title.split('[')[1].split(']')[0]}]")
    ax.set_ylabel("Profondità [m]")
    ax.grid(True, linestyle='--', alpha=0.6)
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf

def crea_report_word_paratia(dati, risultati):
    """Crea un documento Word con i risultati dell'analisi della paratia."""
    doc = Document()
    doc.add_heading('Relazione di Calcolo - Paratia Flessibile', 0)
    doc.add_paragraph(f"Software: ParatieFEM")
    doc.add_paragraph(f"Data: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading('1. Dati di Input', level=1)
    doc.add_paragraph(f"Altezza scavo: {dati.altezza_scavo:.2f} m")
    doc.add_paragraph(f"Lunghezza paratia: {dati.lunghezza_paratia:.2f} m")
    doc.add_paragraph(f"Spessore paratia: {dati.spessore_paratia:.2f} m")
    doc.add_paragraph(f"Sovraccarico: {dati.q_sovraccarico_kPa:.1f} kPa")
    doc.add_paragraph(f"Profondità falda a monte: {dati.falda_monte_m:.2f} m")
    doc.add_paragraph(f"Profondità falda a valle: {dati.falda_valle_m:.2f} m")

    if dati.kh > 0:
        doc.add_heading('Dati Sismici', level=2)
        doc.add_paragraph(f"Coefficiente sismico orizzontale kh: {dati.kh:.3f}")
        doc.add_paragraph(f"Coefficiente sismico verticale kv: {dati.kv:.3f}")

    if not dati.fasi_costruttive_df.empty:
        doc.add_heading('Fasi Costruttive', level=2)
        fasi_df_renamed = dati.fasi_costruttive_df.rename(columns={
            'livello_scavo_m': 'Livello Scavo [m]',
            'descrizione': 'Descrizione'
        })
        add_df_to_doc(doc, fasi_df_renamed)

    if not dati.tiranti_df.empty:
        doc.add_heading('Dati Tiranti / Puntoni', level=2)
        tiranti_df_renamed = dati.tiranti_df.rename(columns={
            'profondita_m': 'Profondità [m]',
            'rigidezza_kN_m': 'Rigidezza [kN/m]',
            'precarico_kN': 'Precarico [kN]',
            'fase_attivazione': 'Fase Attivazione'
        })
        # Assicurati che la colonna fase sia un intero
        tiranti_df_renamed['Fase Attivazione'] = tiranti_df_renamed['Fase Attivazione'].astype(int)
        add_df_to_doc(doc, tiranti_df_renamed)

    doc.add_heading('2. Risultati di Sintesi', level=1)
    doc.add_paragraph(f"Infissione minima richiesta: {risultati['infissione_minima_m']:.2f} m")
    doc.add_paragraph(f"Spostamento massimo in testa: {risultati['spostamento_mm'].max():.1f} mm")
    doc.add_paragraph(f"Momento flettente massimo: {risultati['momento_kNm'].max():.1f} kNm/m")
    fs_globale = risultati.get('fs_globale', -1.0)
    if fs_globale > 0:
        doc.add_paragraph(f"Fattore di Sicurezza (Stabilità Globale, Fellenius): {fs_globale:.2f}")

    doc.add_heading('3. Diagrammi dei Risultati', level=1)
    doc.add_paragraph("Diagramma degli spostamenti:")
    doc.add_picture(plot_risultati_mpl(risultati, 'spostamento_mm', 'Spostamento [mm]'), width=Inches(4.0))

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io.getvalue()